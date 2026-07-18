"""Read a DOCX into the IR.

Headings start slides, the paragraphs, lists and tables until the next heading
are the slide content. Body order is kept by walking the document body. A
numbered-list style becomes an ordered list, monospace runs become code.
"""

from __future__ import annotations

import hashlib
import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .ir import Block, Deck, Image, Segment, Slide

_MONO_RE = re.compile(r"(mono|courier|consol|menlo|monaco|inconsolata)", re.I)

def _is_mono(name: str | None) -> bool:
    return bool(name and _MONO_RE.search(name))

def _para_segments(p: Paragraph) -> tuple[list[Segment], bool]:
    segs: list[Segment] = []
    mono_flags: list[bool] = []
    for run in p.runs:
        if not run.text:
            continue
        mono = _is_mono(run.font.name)
        emph = bool(run.font.bold or run.font.italic)
        segs.append(Segment(text=run.text, kind="code" if mono else "prose", emphasis=emph))
        if run.text.strip():
            mono_flags.append(mono)
    return segs, bool(mono_flags) and all(mono_flags)

def _style(p: Paragraph) -> str:
    try:
        return (p.style.name or "").lower()
    except Exception:
        return ""

def _images(doc) -> dict[str, Image]:
    out: dict[str, Image] = {}
    seen: dict[str, str] = {}
    n = 0
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
        except Exception:
            continue
        h = hashlib.sha1(blob).hexdigest()
        if h in seen:
            continue
        n += 1
        ext = (rel.target_ref.rsplit(".", 1)[-1] or "png").split("/")[-1]
        fname = f"docx_img_{n}.{ext}"
        seen[h] = fname
        out[rel.rId] = Image(filename=fname, data=blob)
    return out

def _open(path: str):
    """Open a .docx, repairing missing [Content_Types] defaults if needed.

    Some documents ship a media part with an odd extension (e.g. `.undefined`)
    that has no declared content type; python-docx then refuses to open the file
    at all. Patch in octet-stream defaults for any missing extension and retry.
    """
    try:
        return Document(path)
    except Exception:
        import io
        import zipfile
        with open(path, "rb") as fh:
            zin = zipfile.ZipFile(io.BytesIO(fh.read()))
        names = zin.namelist()
        ct = "[Content_Types].xml"
        xml = zin.read(ct).decode("utf-8")
        have = set(re.findall(r'Extension="([^"]+)"', xml))
        exts = {n.rsplit(".", 1)[-1].lower() for n in names if "." in n and "/" in n}
        add = "".join(f'<Default Extension="{e}" ContentType="application/octet-stream"/>'
                      for e in sorted(exts) if e not in have)
        xml = xml.replace("</Types>", add + "</Types>")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zo:
            for n in names:
                zo.writestr(n, xml.encode("utf-8") if n == ct else zin.read(n))
        buf.seek(0)
        return Document(buf)

def extract(path: str) -> Deck:
    doc = _open(path)
    deck = Deck(title="", source_lang="en")
    rid_images = _images(doc)

    slides: list[Slide] = []
    cur: Slide | None = None
    list_buf: list[Segment] = []
    list_ordered = False

    def ensure_slide(title: str = "") -> Slide:
        nonlocal cur
        if cur is None:
            cur = Slide(index=len(slides), title=title)
            slides.append(cur)
        return cur

    def flush_list():
        nonlocal list_buf, list_ordered
        if list_buf:
            ensure_slide().blocks.append(
                Block(kind="olist" if list_ordered else "list", segments=list_buf, ordered=list_ordered))
            list_buf = []
            list_ordered = False

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            text = p.text.strip()
            style = _style(p)
            if not text and not p.runs:
                continue
            if style.startswith("heading") or style == "title":
                flush_list()
                cur = Slide(index=len(slides), title=text)
                slides.append(cur)
                continue
            if not text:
                continue
            if "list" in style:
                list_buf.append(Segment(text=text, kind="prose"))
                if "number" in style:
                    list_ordered = True
                continue
            flush_list()
            segs, is_code = _para_segments(p)
            if is_code:
                ensure_slide().blocks.append(
                    Block(kind="code", segments=[Segment(text=p.text, kind="code")]))
            else:
                ensure_slide().blocks.append(Block(kind="para", segments=segs))
        elif child.tag == qn("w:tbl"):
            flush_list()
            table = Table(child, doc)
            rows: list[list[list[Segment]]] = []
            for r in table.rows:
                row = []
                for cell in r.cells:
                    cs: list[Segment] = []
                    for cp in cell.paragraphs:
                        s, _ = _para_segments(cp)
                        cs.extend(s)
                    row.append(cs or [Segment(text="", kind="prose")])
                rows.append(row)
            if rows:
                ensure_slide().blocks.append(Block(kind="table", rows=rows))
    flush_list()

    if rid_images:
        target = slides[0] if slides else ensure_slide("Document")
        target.images.extend(rid_images.values())

    deck.slides = [s for s in slides if s.title or s.blocks or s.images]
    return deck
